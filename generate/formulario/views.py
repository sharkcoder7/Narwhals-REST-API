# -*- coding: UTF-8 -*-

import string
import unicodedata
from StringIO import *

from django.http import HttpResponse, HttpResponseNotFound
from django.template import RequestContext, loader
from django.core.paginator import Paginator
from django.shortcuts import render

# Create your views here.

from formulario.forms import InformeForm
from django.views.generic.edit import FormView

class InformeView(FormView):
    template_name = 'formulario/index.html'
    form_class = InformeForm

    def form_valid(self, form):
        return super(InformeView, self).form_valid(form)


def formulario_save(request):

    from docx import *
    from docx.shared import *

    empresa = request.POST['empresa']
    accionistas_socios = request.POST['accionistas_socios'].lower()
    fecha = request.POST['fecha']
    salvedades = request.POST.get('salvedades', False)
    enfasis = request.POST.get('enfasis', False)
	
    document = Document()
    docx_title="Informe_Auditoria.docx"

    from docx.enum.text import WD_LINE_SPACING

    paragraph = document.add_paragraph().add_run(u'INFORME DE AUDITORÍA INDEPENDIENTE DE CUENTAS ANUALES')

    paragraph.bold = True

    paragraph.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    document.add_paragraph().add_run(u'A los %s de %s. [por encargo de.....]' % (accionistas_socios, empresa)).line_spacing_rule = WD_LINE_SPACING.DOUBLE

    paragraph = document.add_paragraph().add_run(u"Informe sobre las cuentas anuales")

    paragraph.bold = True

    document.add_paragraph(u"Hemos auditado las cuentas anuales adjuntas de la sociedad %s, que comprenden el balance a %s, la cuenta de pérdidas y ganancias, el estado de cambios en el patrimonio neto y la memoria correspondientes al ejercicio terminado en dicha fecha." % (empresa, fecha))

    paragraph = document.add_paragraph().add_run(u"Responsabilidad de los administradores en relación con las cuentas anuales")

    paragraph.italic = True

    paragraph.bold = True

    document.add_paragraph(u"Los administradores son responsables de formular las cuentas anuales adjuntas, de forma que expresen la imagen fiel del patrimonio, de la situación financiera y de los resultados de %s, de conformidad con el marco normativo de información financiera aplicable a la entidad en España, que se identifica en la nota X de la memoria adjunta, y del control interno que consideren necesario para permitir la preparación de cuentas anuales libres de incorrección material, debida a fraude o error." % empresa)

    paragraph = document.add_paragraph().add_run(u"Responsabilidad del auditor")

    paragraph.italic = True

    paragraph.bold = True

    document.add_paragraph(u"Nuestra responsabilidad es expresar una opinión sobre las cuentas anuales adjuntas basada en nuestra auditoría.")

    document.add_paragraph(u"Hemos llevado a cabo nuestra auditoría de conformidad con la normativa reguladora de la auditoría de cuentas vigente en España.")

    document.add_paragraph(u"Dicha normativa exige que cumplamos los requerimientos de ética, así como que planifiquemos y ejecutemos la auditoría con el fin de obtener una seguridad razonable de que las cuentas anuales están libres de incorrecciones materiales.")

    document.add_paragraph(u"Una auditoría requiere la aplicación de procedimientos para obtener evidencia de auditoría sobre los importes y la información revelada en las cuentas anuales.")

    document.add_paragraph(u"Los procedimientos seleccionados dependen del juicio del auditor, incluida la valoración de los riesgos de incorrección material en las cuentas anuales, debida a fraude o error.")

    document.add_paragraph(u"Al efectuar dichas valoraciones del riesgo, el auditor tiene en cuenta el control interno relevante para la formulación por parte de la entidad de las cuentas anuales, con el fin de diseñar los procedimientos de auditoría que sean adecuados en función de las circunstancias, y no con la finalidad de expresar una opinión sobre la eficacia del control interno de la entidad.")

    document.add_paragraph(u"Una auditoría también incluye la evaluación de la adecuación de las políticas contables aplicadas y de la razonabilidad de las estimaciones contables realizadas por la dirección, así como la evaluación de la presentación de las cuentas anuales tomadas en su conjunto.")

    document.add_paragraph(u"Consideramos que la evidencia de auditoría que hemos obtenido proporciona una base suficiente y adecuada para nuestra opinión de auditoría con salvedades.")

    document.add_paragraph()

    document.add_paragraph(u"Fundamento de la opinión con salvedades")

    document.add_paragraph(u"Los valores negociables a corto plazo de la sociedad están valorados en el balance en xxx.")
    document.add_paragraph(u"Los administradores no han actualizado estos valores a valor de mercado sino que, en su lugar, los han registrado al coste, lo que constituye un incumplimiento del marco normativo de información financiera aplicable.")

    document.add_paragraph(u"Los registros de la sociedad indican que, si se hubieran actualizado los valores negociables a valor de mercado, la sociedad habría reconocido unas pérdidas no realizadas de xxx en la cuenta de pérdidas y ganancias del ejercicio.")

    document.add_paragraph(u"El valor registrado de los valores negociables en el balance se habría reducido por el mismo importe a %s, y el impuesto sobre beneficios, el resultado neto y el patrimonio neto se habrían reducido en xxx, xxx, y xxx, respectivamente." % fecha)

    if salvedades: 
        document.add_paragraph()

        document.add_paragraph(u"Opinión con salvedades")

        document.add_paragraph()

        document.add_paragraph(u"En nuestra opinión, excepto por los efectos del hecho descrito en el párrafo de “Fundamento de la opinión con salvedades”, las cuentas anuales adjuntas expresan, en todos los aspectos significativos, la imagen fiel del patrimonio y de la situación financiera de la sociedad %s a %s, así como de sus resultados correspondientes al ejercicio terminado en dicha fecha de conformidad con el marco normativo de información financiera que resulta de aplicación y, en particular, con los principios y criterios contables contenidos en el mismo." % (empresa, fecha))


    if enfasis:
        document.add_paragraph()

        document.add_paragraph(u"Párrafo de énfasis")

        document.add_paragraph()

        document.add_paragraph(u"Llamamos la atención sobre la nota X de las cuentas anuales, que describe una incertidumbre relacionada con el resultado de un litigio emprendido contra la sociedad por la sociedad XYZ. Esta cuestión no modifica nuestra opinión.")


    document.add_page_break()

    # Prepare document for download        
    # -----------------------------
    f = StringIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response


index = InformeView.as_view()

